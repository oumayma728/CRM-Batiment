"""
generer_devis_complet.py — Pipeline complet : description libre (texte) ->
ToolRegistry -> matching catalogue -> document devis final (.docx), fidele
au format reel fourni par l'utilisateur (DEVIS_N__DEV.docx).

Flux :
  1. ToolRegistry.generate_devis_from_text() -- generation libre (trace
     conservee pour audit, mais PAS utilisee pour les prix -- cf. decision
     de conception deja prise dans catalogue_matching.py : prix toujours
     calcules depuis le vrai catalogue, jamais inventes par le modele)
  2. matcher_description_vers_catalogue() -- correspondance SKU + prix reels
  3. Rendu .docx fidele au format reel (en-tete, tableau lignes, totaux
     HT/TVA/TTC, delais, conditions de paiement, garantie, signature)

Point d'attention TVA : taux par defaut 10% (coherent avec l'exemple fourni,
taux reduit renovation FR pour logements >2 ans) -- A CONFIRMER selon le
type reel de travaux (10% renovation, 20% neuf/agrandissement/extension).
"""

from __future__ import annotations

import random
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt

from toolregistry.registry import ToolRegistry
from catalogue_matching import matcher_description_vers_catalogue, LigneMatching


@dataclass
class EntrepriseInfo:
    nom: str
    adresse: str
    tel: str
    email: str


@dataclass
class ClientInfo:
    nom: str
    adresse: str


def generer_numero_devis() -> str:
    from datetime import datetime
    return f"DEV-{datetime.now().year}-{random.randint(10000, 99999):05d}"


def _formater_quantite(quantite: float) -> str:
    """Affiche 80 au lieu de 80.0 quand la quantite est un nombre entier -- plus lisible sur un devis client."""
    if quantite == int(quantite):
        return str(int(quantite))
    return str(quantite)


def generer_devis_depuis_lignes(
    lignes: list[dict],
    entreprise: EntrepriseInfo,
    client: ClientInfo,
    objet: str,
    tva_pct: float = 10.0,
    output_path: Optional[Path] = None,
) -> dict:
    """
    Genere le devis directement depuis des lignes DEJA VALIDEES (potentiellement
    editees par l'utilisateur : quantite, prix unitaire) -- AUCUN appel au
    modele ici, purement deterministe. Utilise par l'ecran Edit Devis pour
    que les corrections manuelles de l'utilisateur soient reellement prises
    en compte, contrairement a generer_devis_complet() qui refait tout le
    pipeline (et donc perd toute modification manuelle) a partir du texte brut.

    lignes : liste de dicts avec au moins label_prestation, quantite_estimee,
    unite, prix_unitaire, sous_total (sous_total recalcule ici de toute
    facon, au cas ou quantite/prix auraient ete modifies sans recalcul cote
    frontend -- jamais fait confiance a un sous_total envoye tel quel).
    """
    lignes_normalisees = []
    total_ht = 0.0
    for l in lignes:
        quantite = float(l.get("quantite_estimee", 0) or 0)
        prix_unitaire = float(l.get("prix_unitaire", 0) or 0)
        sous_total = round(quantite * prix_unitaire, 2)
        total_ht += sous_total
        lignes_normalisees.append(LigneMatching(
            sku_catalogue=l.get("sku_catalogue"),
            label_prestation=l.get("label_prestation", ""),
            quantite_estimee=quantite,
            unite=l.get("unite", ""),
            prix_unitaire=prix_unitaire,
            sous_total=sous_total,
            matching_confidence=float(l.get("matching_confidence", 1.0)),
            non_trouve_dans_catalogue=False,
        ))

    total_ht = round(total_ht, 2)
    montant_tva = round(total_ht * tva_pct / 100, 2)
    total_ttc = round(total_ht + montant_tva, 2)

    resultat = {
        "numero_devis": generer_numero_devis(),
        "date": date.today().strftime("%d/%m/%Y"),
        "validite_jours": 30,
        "entreprise": entreprise,
        "client": client,
        "objet": objet,
        "modalite_source": "edition_manuelle",
        "lignes_trouvees": lignes_normalisees,
        "lignes_non_trouvees": [],
        "total_ht": total_ht,
        "tva_pct": tva_pct,
        "montant_tva": montant_tva,
        "total_ttc": total_ttc,
        "prestations_generees_ia": None,
    }

    if output_path:
        rendre_docx(resultat, output_path)

    return resultat


def generer_devis_complet(
    description_libre: str,
    modalite_source: str,
    entreprise: EntrepriseInfo,
    client: ClientInfo,
    tva_pct: float = 10.0,
    output_path: Optional[Path] = None,
) -> dict:
    registry = ToolRegistry()

    # Etape 1 : generation libre (trace conservee pour audit uniquement)
    devis_libre = registry.generate_devis_from_text(description_libre)

    # Etape 2 : matching catalogue -- prix TOUJOURS calcules depuis le vrai
    # catalogue, jamais depuis la generation libre du modele.
    lignes_matching = matcher_description_vers_catalogue(description_libre, registry=registry)

    lignes_trouvees = [l for l in lignes_matching if not l.non_trouve_dans_catalogue]
    lignes_non_trouvees = [l for l in lignes_matching if l.non_trouve_dans_catalogue]

    total_ht = round(sum(l.sous_total for l in lignes_trouvees), 2) if lignes_trouvees else 0.0
    montant_tva = round(total_ht * tva_pct / 100, 2)
    total_ttc = round(total_ht + montant_tva, 2)

    resultat = {
        "numero_devis": generer_numero_devis(),
        "date": date.today().strftime("%d/%m/%Y"),
        "validite_jours": 30,
        "entreprise": entreprise,
        "client": client,
        "objet": description_libre,
        "modalite_source": modalite_source,
        "lignes_trouvees": lignes_trouvees,
        "lignes_non_trouvees": lignes_non_trouvees,
        "total_ht": total_ht,
        "tva_pct": tva_pct,
        "montant_tva": montant_tva,
        "total_ttc": total_ttc,
        "prestations_generees_ia": devis_libre.contenu_json,
    }

    if output_path:
        rendre_docx(resultat, output_path)

    return resultat


def rendre_docx(devis: dict, output_path: Path) -> None:
    doc = Document()

    doc.add_heading(f"DEVIS N° {devis['numero_devis']}", level=1)

    table_entete = doc.add_table(rows=1, cols=2)
    ent = devis["entreprise"]
    table_entete.cell(0, 0).text = f"Entreprise\n{ent.nom}\n{ent.adresse}\nTél. : {ent.tel}\nEmail : {ent.email}"
    cli = devis["client"]
    table_entete.cell(0, 1).text = f"Client\n{cli.nom}\n{cli.adresse}"

    doc.add_paragraph(f"Date : {devis['date']}\nValidité : {devis['validite_jours']} jours")

    doc.add_heading("Objet", level=2)
    doc.add_paragraph(devis["objet"])

    doc.add_heading("Récapitulatif", level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Désignation", "Qté", "Unité", "PU HT (€)", "Total HT (€)"]):
        table.rows[0].cells[i].text = h

    for l in devis["lignes_trouvees"]:
        row = table.add_row().cells
        row[0].text = l.label_prestation
        row[1].text = _formater_quantite(l.quantite_estimee)
        row[2].text = l.unite
        row[3].text = f"{l.prix_unitaire:.2f}"
        row[4].text = f"{l.sous_total:.2f}"

    if devis["lignes_non_trouvees"]:
        doc.add_heading("Prestations complémentaires (hors catalogue, à chiffrer manuellement)", level=2)
        for l in devis["lignes_non_trouvees"]:
            doc.add_paragraph(l.label_prestation, style="List Bullet")

    table_totaux = doc.add_table(rows=3, cols=2)
    table_totaux.style = "Light Grid Accent 1"
    table_totaux.rows[0].cells[0].text = "Total HT"
    table_totaux.rows[0].cells[1].text = f"{devis['total_ht']:.2f} €"
    table_totaux.rows[1].cells[0].text = f"TVA ({devis['tva_pct']:.0f} %)"
    table_totaux.rows[1].cells[1].text = f"{devis['montant_tva']:.2f} €"
    table_totaux.rows[2].cells[0].text = "Total TTC"
    table_totaux.rows[2].cells[1].text = f"{devis['total_ttc']:.2f} €"

    doc.add_heading("Délais", level=2)
    doc.add_paragraph(devis.get("delai_debut", "Début des travaux : sous 3 semaines après acceptation."))
    doc.add_paragraph(devis.get("delai_duree", "Durée estimée : à préciser selon devis détaillé."))

    doc.add_heading("Conditions de paiement", level=2)
    doc.add_paragraph("Acompte à la commande : 30 %")
    doc.add_paragraph("Début des travaux : 40 %")
    doc.add_paragraph("Solde à la réception : 30 %")

    doc.add_heading("Garantie", level=2)
    doc.add_paragraph("Garantie décennale.")
    doc.add_paragraph("Garantie fabricant sur les matériaux.")
    doc.add_paragraph("Assurance responsabilité civile professionnelle.")

    doc.add_heading("Signature", level=2)
    doc.add_paragraph("Le client reconnaît avoir pris connaissance du présent devis et accepte les travaux décrits ci-dessus.")
    doc.add_paragraph("Bon pour accord")
    doc.add_paragraph("Nom : _______________________")
    doc.add_paragraph("Date : ____ / ____ / ______")
    doc.add_paragraph("Signature : _______________________")

    doc.save(str(output_path))